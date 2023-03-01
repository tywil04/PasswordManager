# This is a really hacky script that scrapes my golang endpoints to generate most of the API documentation. Description and examples need to be added manually.
#
# Please dont use this, its specific to my setup and conventions.

import os
import json
from docx import Document
from docx.shared import Pt
import re

path = "/home/tyler/Development/PasswordManager5/api/endpoints"
apiPrefix = "/api/v1"
dirPathPrefix = "/home/tyler/Development/PasswordManager5/api/endpoints"
outputPath = "/home/tyler/Development/PasswordManager5/docs/api"

errorResponse = {
    "error": {
        "code": "string",
        "message": "string",
    }
}

data = {}

def formatMarkdown(path, entries):
    content = f"""# {path}
If a GET request requires parameters, the parameters can only be passed using query parameters. (GET requests does not support body).

All requests need to have a `Content-type` header that is either `application/json` or `application/xml`, requests wont work otherwise. 

Required/Optional Request/Response parameters are denoted like so:
- `required param` or `<required param>`
- `[optional param]` 

This API is public, however, it is strongly recommended that you use an official client."""
    
    for method, data in entries.items():
        content += f"""

## {method}
### Description
// DESCRIPTION HERE //

### Request Format
#### Headers
```json
{json.dumps(data["headers"], indent=4) or "None"}
```

#### Params
```json
{json.dumps(data["input"], indent=4) or "None"}
```

### Response Format
"""

        for status, responseData in sorted(data["responses"].items(), key=lambda x: int(x[0])):
            if status == 500:
                status = "500 (Server Error)"
            elif status == 400:
                status = "400 (Client Error)"
                
            content += f"""#### {status}
```json
{json.dumps(responseData, indent=4) or "None"}
```
"""
        content += f"""
### Example
```javascript
// JAVASCRIPT EXAMPLE HERE
```"""

    return content

def startDocxFormat(doc):
    doc.add_page_break()
    doc.add_heading("API Docs", level=1)
    doc.add_paragraph("If a GET request requires parameters, the parameters can only be passed using query parameters. (GET requests does not support body).")
    doc.add_paragraph("All requests need to have a `Content-type` header that is either `application/json` or `application/xml`, requests wont work otherwise.")
    doc.add_paragraph("Required/Optional Request/Response parameters are denoted like so:")
    doc.add_paragraph("`required param` or `<required param>`", style="List Paragraph")
    doc.add_paragraph("`[optional param]` ", style="List Paragraph")
    doc.add_paragraph("This API is public, however, it is strongly recommended that you use an official client.")

def formatDocx(doc, path, entries):
    doc.add_page_break()

    style = doc.styles['Normal']
    font = style.font
    font.name = "Calibri (Body)"
    font.size = Pt(11)

    tableStyle = doc.styles["List Table 1 Light Accent 5"]
    tableFont = tableStyle.font
    tableFont.name = "Cascadia Code"
    tableFont.size = Pt(8)
    tableFont.bold = False

    doc.add_heading(path, level=1)
    
    for method, data in entries.items():
        doc.add_heading(method, level=2)
        
        doc.add_heading("Description", level=3)

        if "description" in list(data.keys()):
            doc.add_paragraph(data["description"], style=style)
        else:
            doc.add_paragraph("// DESCRIPTION HERE //", style=style)

        # doc.add_paragraph("", style="No Spacing")
        doc.add_heading("Request Format", level=3)
        
        table = doc.add_table(rows=2, cols=2, style=tableStyle)
        table.rows[0].cells[0].text = "Headers"
        for paragraph in table.rows[0].cells[0].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True 

        table.rows[0].cells[1].text = "Params"
        for paragraph in table.rows[0].cells[1].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True 

        table.rows[1].cells[0].text = json.dumps(data["headers"], indent=4).replace('": "', '": ').replace('",\n', ",\n").replace('"\n', "\n").replace('"string', "string")

        if table.rows[1].cells[0].text == "{}":
            table.rows[1].cells[0].text = "None"

        for paragraph in table.rows[1].cells[0].paragraphs:
            for run in paragraph.runs:
                run.font.name = "Cascadia Code"
                run.font.size = Pt(8)
                run.font.bold = False 

        table.rows[1].cells[1].text = json.dumps(data["input"], indent=4).replace('": "', '": ').replace('",\n', ",\n").replace('"\n', "\n").replace('"string', "string")

        if table.rows[1].cells[1].text == "{}":
            table.rows[1].cells[1].text = "None"

        for paragraph in table.rows[1].cells[1].paragraphs:
            for run in paragraph.runs:
                run.font.name = "Cascadia Code"
                run.font.size = Pt(8)
                run.font.bold = False 

        # doc.add_paragraph("", style="No Spacing")
        doc.add_heading("Response Format", level=3)

        table1 = doc.add_table(rows=2, cols=1, style=tableStyle)
        table1.autofit = True
        table1.allow_autofit = True

        if len(data["responses"]) - 1 != 0:
            table2 = doc.add_table(rows=2, cols=len(data["responses"]) - 1, style=tableStyle)
            table2.autofit = True
            table2.allow_autofit = True

            table1Selected = True
            col = 0

            for status, responseData in sorted(data["responses"].items(), key=lambda x: int(x[0])):
                if table1Selected:
                    table1.rows[0].cells[col].text = f"Status {status}"
                    for paragraph in table1.rows[0].cells[col].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True 

                    table1.rows[1].cells[col].text = json.dumps(responseData, indent=4).replace('": "', '": ').replace('",\n', ",\n").replace('"\n', "\n").replace('"string', "string")

                    if table1.rows[1].cells[col].text == "{}":
                        table1.rows[1].cells[col].text = "None"

                    for paragraph in table1.rows[1].cells[col].paragraphs:
                        for run in paragraph.runs:
                            run.font.name = "Cascadia Code"
                            run.font.size = Pt(8)
                            run.font.bold = False 
                    table1Selected = not table1Selected
                else:
                    table2.rows[0].cells[col].text = f"Status {status}"
                    for paragraph in table2.rows[0].cells[col].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True 

                    table2.rows[1].cells[col].text = json.dumps(responseData, indent=4).replace('": "', '": ').replace('",\n', ",\n").replace('"\n', "\n").replace('"string', "string")

                    if table2.rows[1].cells[col].text == "{}":
                        table2.rows[1].cells[col].text = "None"

                    for paragraph in table2.rows[1].cells[col].paragraphs:
                        for run in paragraph.runs:
                            run.font.name = "Cascadia Code"
                            run.font.size = Pt(8)
                            run.font.bold = False 
                    col += 1

        # doc.add_paragraph("", style="No Spacing")
        doc.add_heading("Example", level=3)
        doc.add_paragraph("// JAVASCRIPT EXAMPLE HERE //")

for (dirPath, dirNames, fileNames) in os.walk(path):
    for fileName in fileNames:
        endpointPath = apiPrefix + (dirPath + "/" + fileName).replace(dirPathPrefix, "")

        splitEndpointPath = endpointPath.split("/")
        folderName = splitEndpointPath[-2:-1][0]
        goFileName = splitEndpointPath[-1:][0].replace(".go", "")

        if folderName == goFileName:
            endpointPath = endpointPath.replace("/" + folderName + ".go", "")
        else:
            endpointPath = endpointPath.replace(".go", "")

        with open(dirPath + "/" + fileName, "r") as fileReader:
            data[endpointPath] = {}
            dataIndex = -1
            dataKey = ""
            descs = []

            finishedStructs = []
            structDataLocation = {}
            structContinue = False
            embededStructDataLocation = {}
            embededStructContinue = False
            embededIsArray = False
            thirdEmbededStructDataLocation = {}
            thirdEmbededStructContinue = False
            thirdEmbededIsArray = False

            for line in fileReader.readlines():
                if line.startswith("func"):
                    try:
                        data[endpointPath][dataKey]["description"] = descs[dataIndex][:-1]
                    except: pass
                    dataIndex += 1
                    dataKey = list(data[endpointPath].keys())[dataIndex]
                    
                if line.startswith("type ") and (line.endswith(" struct{}\n") or line.endswith(" struct {\n")):
                    structName = line.replace("type ", "").replace(" struct{}\n", "").replace(" struct {\n", "")

                    method = ""
                    index = 0
                    while structName[index].islower() or index == 0:
                        method += structName[index]
                        index += 1

                    method = method.upper()

                    data[endpointPath][method] = {
                        "input": {},
                        "responses": {},
                        "headers": {},
                    }

                if "Input" in line and line.endswith(" struct {\n"):
                    structContinue = True
                elif "Input" in line and line.endswith(" struct{}\n"):
                    finishedStructs.append({})

                if "c.JSON(400, " in line:
                    data[endpointPath][dataKey]["responses"][400] = errorResponse
                    
                if "c.JSON(500, " in line:
                    data[endpointPath][dataKey]["responses"][500] = errorResponse

                if "authedUser" in line:
                    data[endpointPath][dataKey]["headers"]["Authorization"] = "valid authToken"
                        
                if "c.JSON(200, " in line:
                    response = "{" + line.strip().replace("c.JSON(200, gin.H{", "").replace("})", "") + "}"
                    jsonResponse = response.replace(": ", ': "').replace(", ", '", ').replace("}", '"}').replace('""', '"').replace('" + ', " ").replace('{"}', "{}")
                    responseDict = json.loads(jsonResponse)
                    returnResponseDict = {}

                    for (key, value) in responseDict.items():
                        splitValue = value.split(".")
                        if len(splitValue) == 3 and splitValue[-2:-1][0] == "ID" and splitValue[-1:][0] == "String()":
                            returnResponseDict[key] = "uuid string"
                        elif key == "options":
                            returnResponseDict[key] = {
                                "publicKey": {
                                    "challenge": "url base64 string",
                                    "rp": {
                                        "name": "string",
                                        "[icon]": "string",
                                        "id": "string",
                                    },
                                    "user": {
                                        "name": "string",
                                        "[icon]": "string",
                                        "[displayName]": "string",
                                        "id": "url base64 string",
                                    },
                                    "[pubKeyCredParams]": [
                                        {
                                            "type": "string",
                                            "alg": "int",
                                        }
                                    ],
                                    "[authenticatorSelection]": {
                                        "[authenticatorAttachment]": "string",
                                        "[requireResidentKey]": "bool",
                                        "[residentKey]": "string",
                                        "[userVerification]": "string",
                                    },
                                    "[timeout]": "int",
                                    "[excludeCredentials]": [
                                        {
                                            "type": "string",
                                            "id": "url base64 string",
                                            "[transports]": ["string"]
                                        }
                                    ],
                                    "[attestation]": "string",
                                }
                            }
                        elif key == "availableChallenges":
                            returnResponseDict[key] = ["string"]
                        elif key == "webauthnCredentials":
                            returnResponseDict[key] = [
                                {
                                    "id": "uuid string",
                                    "name": "string",
                                    "createdAt": "time string",
                                }
                            ]
                        elif key == "passwords":
                            returnResponseDict[key] = [
                                {
                                    "id": "uuid string",
                                    "name": "base64 string",
                                    "nameIv": "base64 string",
                                    "username": "base64 string",
                                    "usernameIv": "base64 string",
                                    "password": "base64 string",
                                    "passwordIv": "base64 string",
                                    "colour": "hex colour string",
                                    "additionalFields": [
                                        {
                                            "key": "base64 string",
                                            "keyIv": "base64 string",
                                            "value": "base64 string",
                                            "valueIv": "base64 string",
                                        },
                                    ],
                                    "urls": [
                                        {
                                            "url": "base64 string",
                                            "urlIv": "base64 string",
                                        },
                                    ],
                                }
                            ]
                        elif key == "authToken":
                            returnResponseDict[key] = "authToken string"
                        else:
                            returnResponseDict[key] = "base64 string"

                    data[endpointPath][dataKey]["responses"][200] = returnResponseDict
                    if len(finishedStructs) > dataIndex:
                        data[endpointPath][dataKey]["input"] = finishedStructs[dataIndex]

                if bool(re.search(r".+Description string = \"", line)):
                    descs.append(re.sub(r".+Description string \"", "", line)[:-1].replace("\t", "").replace('"', ""))

                if thirdEmbededStructContinue:
                    stripped = line.strip()
                    if stripped == "}":
                        thirdEmbededStructContinue = False
                    else:
                        parts = line.strip().split(" `")
                        if parts[0] == "}":
                            thirdEmbededStructContinue = False
                            name =  parts[1].replace('form:"', "").replace('" json:"', "//").split("//")[0]
                            if thirdEmbededIsArray:
                                embededStructDataLocation[name] = [thirdEmbededStructDataLocation]
                            else:
                                embededStructDataLocation[name] = thirdEmbededStructDataLocation
                            thirdEmbededIsArray = False
                            continue
                        elif len(parts) == 2:
                            expectedType = parts[0].replace("  ", " ").replace("  ", " ").replace("  ", " ").replace("  ", " ").replace("  ", " ").replace("  ", " ").replace("  ", " ").replace("  ", " ").split(" ")[1]
                            name = parts[1].replace('form:"', "").replace('" json:"', "//").split("//")[0]
                            thirdEmbededStructDataLocation[name] = expectedType

                if embededStructContinue and not thirdEmbededStructContinue:
                    stripped = line.strip()
                    if stripped == "}":
                        embededStructContinue = False
                    else:
                        parts = line.strip().split(" `")
                        if parts[0] == "}":
                            embededStructContinue = False
                            name =  parts[1].replace('form:"', "").replace('" json:"', "//").split("//")[0]
                            if embededIsArray:
                                structDataLocation[name] = [embededStructDataLocation]
                            else:
                                structDataLocation[name] = embededStructDataLocation
                            embededIsArray = False
                            continue
                        else:
                            expectedType = parts[0].replace("  ", " ").replace("  ", " ").replace("  ", " ").replace("  ", " ").replace("  ", " ").replace("  ", " ").replace("  ", " ").replace("  ", " ").split(" ")[1]
                            if expectedType == "struct":
                                thirdEmbededStructContinue = True
                                thirdEmbededStructDataLocation = {}
                            elif expectedType == "[]struct":
                                thirdEmbededStructContinue = True
                                thirdEmbededStructDataLocation = {}
                                thirdEmbededIsArray = True
                            elif len(parts) == 2:
                                name = parts[1].replace('form:"', "").replace('" json:"', "//").split("//")[0]
                                embededStructDataLocation[name] = expectedType

                if structContinue and not embededStructContinue and not thirdEmbededStructContinue:
                    stripped = line.strip()
                    if stripped == "}":
                        structContinue = False
                        finishedStructs.append(structDataLocation)
                    else:
                        parts = line.strip().split(" `")
                        if parts[0] == "}":
                            embededStructContinue = False
                        else:
                            expectedType = parts[0].replace("  ", " ").replace("  ", " ").replace("  ", " ").replace("  ", " ").replace("  ", " ").replace("  ", " ").replace("  ", " ").replace("  ", " ").split(" ")[1]
                            if expectedType == "struct":
                                embededStructContinue = True
                                embededStructDataLocation = {}
                            elif expectedType == "[]struct":
                                embededStructContinue = True
                                embededStructDataLocation = {}
                                embededIsArray = True
                            elif len(parts) == 2:
                                name = parts[1].replace('form:"', "").replace('" json:"', "//").split("//")[0]
                                structDataLocation[name] = expectedType

doc = Document(docx="/home/tyler/Development/PasswordManager5/docs/Tyler Williams - Computer Science Project.docx")
startDocxFormat(doc)
for path, data in data.items():
    # markdown = formatMarkdown(path, data)
    # fileName = f"{outputPath}{path.replace(apiPrefix, '')}.md"
    # os.makedirs(os.path.dirname(fileName), exist_ok=True)
    # with open(fileName, "w+") as fileWriter:
    #     fileWriter.write(markdown)
    formatDocx(doc, path, data)

doc.save(outputPath + "/test.docx")

print("Success")
